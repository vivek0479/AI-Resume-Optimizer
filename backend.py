import re
import os
import time
import io
import json

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx as python_docx
except ImportError:
    python_docx = None

# ─── API key handling (environment / Streamlit secrets / legacy config) ───
_MODEL_NAME = "gemini-2.5-flash"
_CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")
_STREAMLIT_SECRETS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".streamlit", "secrets.toml")


def _load_streamlit_secrets() -> str:
    """Read Gemini API key from Streamlit secrets if available."""
    try:
        import streamlit as st
        secrets = getattr(st, "secrets", None)
        if secrets:
            value = secrets.get("GEMINI_API_KEY", "")
            if isinstance(value, str):
                return value.strip()
    except Exception:
        pass
    return ""


def _load_secrets_file() -> str:
    """Read Gemini API key from a local .streamlit/secrets.toml file."""
    if not os.path.exists(_STREAMLIT_SECRETS_PATH):
        return ""

    try:
        with open(_STREAMLIT_SECRETS_PATH, "r", encoding="utf-8") as f:
            content = f.read()

        try:
            import tomllib
            data = tomllib.loads(content)
        except Exception:
            data = {}
            section = None
            for line in content.splitlines():
                cleaned = line.strip()
                if not cleaned or cleaned.startswith("#"):
                    continue
                if cleaned.startswith("[") and cleaned.endswith("]"):
                    section = cleaned[1:-1].strip().lower()
                    continue
                if "=" in cleaned and section:
                    key, value = [part.strip() for part in cleaned.split("=", 1)]
                    if value.startswith('"') and value.endswith('"'):
                        value = value[1:-1]
                    data.setdefault(section, {})[key] = value

        if isinstance(data, dict):
            for section_name in ("general", "default"):
                section = data.get(section_name)
                if isinstance(section, dict):
                    value = section.get("GEMINI_API_KEY", "")
                    if isinstance(value, str) and value.strip():
                        return value.strip()
            if isinstance(data.get("GEMINI_API_KEY"), str):
                return data["GEMINI_API_KEY"].strip()
    except Exception:
        pass

    return ""


def _get_persisted_api_key() -> str:
    """Return the stored API key from env, Streamlit secrets, or the local secrets file."""
    env_key = os.getenv("GEMINI_API_KEY", "").strip()
    if env_key:
        return env_key

    streamlit_key = _load_streamlit_secrets()
    if streamlit_key:
        return streamlit_key

    return _load_secrets_file()


def save_api_key(api_key: str) -> None:
    """Persist the API key securely in a local Streamlit secrets file."""
    api_key = (api_key or "").strip()
    os.makedirs(os.path.dirname(_STREAMLIT_SECRETS_PATH), exist_ok=True)

    if api_key:
        os.environ["GEMINI_API_KEY"] = api_key
        with open(_STREAMLIT_SECRETS_PATH, "w", encoding="utf-8") as f:
            f.write("[general]\n")
            f.write(f'GEMINI_API_KEY = "{api_key}"\n')
    else:
        os.environ.pop("GEMINI_API_KEY", None)
        if os.path.exists(_STREAMLIT_SECRETS_PATH):
            try:
                os.remove(_STREAMLIT_SECRETS_PATH)
            except Exception:
                pass


def load_config() -> dict:
    """Load the persisted API key from secure storage, migrating legacy config.json if needed."""
    try:
        with open(_CONFIG_PATH, "r", encoding="utf-8") as f:
            legacy_data = json.load(f)
        if isinstance(legacy_data, dict):
            legacy_key = str(legacy_data.get("user_api_key", "")).strip()
            if legacy_key:
                save_api_key(legacy_key)
                try:
                    os.remove(_CONFIG_PATH)
                except Exception:
                    pass
                return {"user_api_key": legacy_key}
    except Exception:
        pass

    return {"user_api_key": _get_persisted_api_key()}


def save_config(data: dict) -> None:
    """Save config data securely without writing plaintext secrets to the repository."""
    if "user_api_key" in data:
        save_api_key(data.get("user_api_key", ""))


def get_active_api_key() -> str:
    """Return the API key to use for Gemini calls."""
    cfg = load_config()
    return cfg.get("user_api_key", "").strip()


def is_using_default_key() -> bool:
    """True when no personal API key is configured."""
    return not bool(get_active_api_key())

def is_quota_error(error: Exception) -> bool:
    """Detect Gemini rate-limit / quota-exceeded errors."""
    msg = str(error).lower()
    quota_keywords = [
        "quota", "rate limit", "resource_exhausted", "429",
        "too many requests", "exceeded", "billing"
    ]
    return any(kw in msg for kw in quota_keywords)


# A standard dictionary of common skills for keyword matching in fallback mode
COMMON_SKILLS = [
    "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Go", "Rust", "Ruby", "PHP", "Swift", "Kotlin",
    "HTML", "CSS", "React", "Angular", "Vue", "Next.js", "Nuxt.js", "Node.js", "Express", "Django", "Flask",
    "FastAPI", "Spring Boot", "ASP.NET", "SQL", "NoSQL", "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch",
    "AWS", "Azure", "GCP", "Docker", "Kubernetes", "CI/CD", "Git", "GitHub", "Terraform", "Ansible", "Linux",
    "Machine Learning", "Deep Learning", "NLP", "Computer Vision", "TensorFlow", "PyTorch", "Scikit-Learn",
    "Data Analysis", "Pandas", "NumPy", "Tableau", "Power BI", "Excel", "R", "SQL Server", "Oracle",
    "Agile", "Scrum", "Kanban", "Project Management", "Product Management", "System Design", "Microservices",
    "REST API", "GraphQL", "gRPC", "WebSockets", "OAuth", "JWT", "Docker Compose", "Nginx", "Apache", "UI/UX",
    "Figma", "Adobe XD", "Photoshop", "Illustrator", "Tailwind CSS", "Bootstrap", "Sass", "Jest", "Cypress",
    "Mocha", "Selenium", "JUnit", "PyTest", "Postman", "Jira", "Confluence", "Trello", "Slack", "VS Code"
]

def parse_resume(file_bytes, filename) -> str:
    """
    Parses resume content from PDF or DOCX bytes.
    """
    ext = os.path.splitext(filename)[1].lower()
    text = ""

    if ext == ".pdf":
        if pypdf:
            try:
                pdf_file = io.BytesIO(file_bytes)
                reader = pypdf.PdfReader(pdf_file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e:
                text = f"[Error reading PDF: {str(e)}]"
        else:
            text = "pypdf not installed. Please run: pip install pypdf"

    elif ext == ".docx":
        if python_docx:
            try:
                doc_file = io.BytesIO(file_bytes)
                doc = python_docx.Document(doc_file)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())
                text = "\n".join(paragraphs)
            except Exception as e:
                text = f"[Error reading DOCX: {str(e)}]"
        else:
            text = "python-docx not installed. Please run: pip install python-docx"
    else:
        # Plain text fallback
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = "Could not parse file."

    return text.strip()

def extract_metadata(text, filename, file_size_bytes):
    """
    Extracts basic info and statistics from the resume text.
    """
    word_count = len(re.findall(r'\w+', text))
    char_count = len(text)
    
    # Estimate pages based on characters (approx 3000 chars per page)
    estimated_pages = max(1, int(char_count / 3000) + (1 if char_count % 3000 > 1000 else 0))
    
    # Format size
    if file_size_bytes < 1024 * 1024:
        size_str = f"{file_size_bytes / 1024:.1f} KB"
    else:
        size_str = f"{file_size_bytes / (1024 * 1024):.2f} MB"
        
    # Detect Skills from common dictionary
    detected_skills = []
    for skill in COMMON_SKILLS:
        # Use regex boundary match for skills
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text, re.IGNORECASE):
            detected_skills.append(skill)
            
    # Fallback if no skills detected
    if not detected_skills:
        detected_skills = ["Communication", "Problem Solving", "Adaptability", "Organization"]
        
    # Try to extract sections (Experience, Education, Projects)
    sections = {
        "Experience": False,
        "Education": False,
        "Projects": False,
        "Skills": len(detected_skills) > 0
    }
    
    if re.search(r'\b(experience|work|employment|history|career)\b', text, re.IGNORECASE):
        sections["Experience"] = True
    if re.search(r'\b(education|university|college|degree|academic)\b', text, re.IGNORECASE):
        sections["Education"] = True
    if re.search(r'\b(project|projects|portfolio)\b', text, re.IGNORECASE):
        sections["Projects"] = True
        
    return {
        "filename": filename,
        "size": size_str,
        "pages": estimated_pages,
        "skills": detected_skills[:12],  # Return top 12 detected skills for preview
        "sections": [k for k, v in sections.items() if v]
    }

def analyze_resume_mock(text, metadata, job_desc=""):
    """
    Simulates a high-fidelity ATS and AI analysis if no Gemini key is set.
    """
    skills = metadata["skills"]
    sections = metadata["sections"]
    
    # Check match with Job Description keywords if provided
    jd_keywords = []
    missing_skills = []
    found_keywords = []
    
    if job_desc:
        # Extract potential keywords from JD
        jd_words = re.findall(r'\b[A-Za-z0-9+#.-]+\b', job_desc)
        jd_candidates = [w for w in jd_words if len(w) > 2 and w[0].isupper()]
        # Unique title-cased candidates
        jd_keywords = sorted(list(set(jd_candidates)))[:10]
        
        # Calculate found and missing
        for kw in jd_keywords:
            if re.search(r'\b' + re.escape(kw) + r'\b', text, re.IGNORECASE):
                found_keywords.append(kw)
            else:
                missing_skills.append(kw)
    else:
        # Fallback keyword list
        found_keywords = skills[:6]
        missing_skills = ["Docker", "Kubernetes", "System Design", "Unit Testing", "CI/CD"]
        
    # Calculate simulated ATS score based on sections found, skills list, and found keywords
    score_base = 40
    if "Experience" in sections: score_base += 15
    if "Education" in sections: score_base += 10
    if "Projects" in sections: score_base += 10
    score_base += min(15, len(skills) * 1.5)
    
    if job_desc:
        keyword_pct = len(found_keywords) / max(1, len(jd_keywords))
        score_base += int(keyword_pct * 10)
    else:
        score_base += 10
        
    ats_score = min(98, max(35, score_base))
    
    # Define color scale for progress elements
    if ats_score >= 80:
        status_color = "#16A34A"  # Success Green
        summary_desc = "Excellent match! Your resume contains strong section alignments, formatting, and key terminology corresponding to industry standards."
    elif ats_score >= 60:
        status_color = "#F59E0B"  # Warning Orange
        summary_desc = "Good foundation, but there's room for optimization. Adding more metrics and matching specific keywords from your target job description will improve your performance."
    else:
        status_color = "#DC2626"  # Error Red
        summary_desc = "Significant improvement needed. Essential sections might be missing, and core industry keywords should be integrated to pass automated screening filters."

    # Analysis Results payload
    return {
        "ats_score": ats_score,
        "status_color": status_color,
        "summary": summary_desc,
        "breakdown": [
            {"category": "Formatting & Structure", "score": 90 if "Experience" in sections else 55, "icon": "📐"},
            {"category": "Keyword Optimization", "score": int(ats_score * 0.9), "icon": "🔑"},
            {"category": "Experience & Impact", "score": 85 if "Experience" in sections else 40, "icon": "💼"},
            {"category": "Skills & Competencies", "score": min(100, len(skills) * 10), "icon": "🎓"}
        ],
        "missing_skills": missing_skills,
        "found_keywords": found_keywords,
        "suggestions": [
            "Action-Oriented Verbs: Replace passive phrases with strong action verbs (e.g., 'Led design', 'Optimized architecture' instead of 'Responsible for...').",
            "Quantify Results: Include measurable achievements where possible (e.g., 'Reduced load times by 40%', 'Managed a team of 4 devs').",
            "Highlight Industry Standard Skills: Integrate missing core skills like " + ", ".join(missing_skills[:3]) + " directly in your experience descriptions.",
            "Formatting Consistency: Ensure all work experiences follow the same layout structure (Title, Company, Date, bullet points)."
        ],
        "recommended_actions": [
            {"action": "Add Metrics & KPIs", "difficulty": "Medium", "impact": "High"},
            {"action": "Tailor to Target Role", "difficulty": "Easy", "impact": "High"},
            {"action": "Restructure Skills Section", "difficulty": "Easy", "impact": "Medium"}
        ]
    }

def analyze_resume_gemini(text, metadata, job_desc="", api_key=""):
    """
    Performs live ATS analysis using the Gemini API.
    """
    import google.generativeai as genai
    import json
    
    genai.configure(api_key=api_key)
    
    # Prompt instructing Gemini to return a clean JSON payload
    prompt = f"""
    You are an expert ATS (Applicant Tracking System) parser and senior recruiter.
    Analyze the following resume text and job description (if provided), and return a strictly formatted JSON report.
    
    RESUME TEXT:
    {text}
    
    JOB DESCRIPTION:
    {job_desc}
    
    Return your response strictly in the following JSON format without any markdown wrappers or explanatory text.
    {{
        "ats_score": (int between 0 and 100),
        "summary": (brief 2-3 sentence overview of the resume alignment and match),
        "breakdown": [
            {{"category": "Formatting & Structure", "score": (int 0-100), "icon": "📐"}},
            {{"category": "Keyword Optimization", "score": (int 0-100), "icon": "🔑"}},
            {{"category": "Experience & Impact", "score": (int 0-100), "icon": "💼"}},
            {{"category": "Skills & Competencies", "score": (int 0-100), "icon": "🎓"}}
        ],
        "missing_skills": [(list of up to 5 critical skills or keywords missing from the resume)],
        "found_keywords": [(list of up to 6 key professional keywords found in the resume)],
        "suggestions": [(list of 4 specific, actionable improvement bullet points)],
        "recommended_actions": [
            {{"action": "Add Metrics & KPIs", "difficulty": "Medium", "impact": "High"}},
            {{"action": "Tailor to Target Role", "difficulty": "Easy", "impact": "High"}},
            {{"action": "Restructure Skills Section", "difficulty": "Easy", "impact": "Medium"}}
        ]
    }}
    """
    
    try:
        model = genai.GenerativeModel(_MODEL_NAME)
        response = model.generate_content(prompt)
        
        # Parse JSON from response text
        response_text = response.text.strip()
        # Remove potential markdown wrappers
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        response_text = response_text.strip()
        
        result = json.loads(response_text)
        
        # Enforce status color color coding
        score = result.get("ats_score", 70)
        if score >= 80:
            result["status_color"] = "#16A34A"
        elif score >= 60:
            result["status_color"] = "#F59E0B"
        else:
            result["status_color"] = "#DC2626"
            
        return result
        
    except Exception as e:
        mock_result = analyze_resume_mock(text, metadata, job_desc)
        if is_quota_error(e):
            mock_result["quota_exceeded"] = True
            mock_result["summary"] = "⚠ The built-in API quota has been reached. Showing simulated results. Please add your own Gemini API key in Settings to continue with real AI analysis."
        else:
            mock_result["summary"] = f"Gemini error: {str(e)[:120]}. Displaying simulated results."
        return mock_result


def validate_api_key(api_key: str) -> dict:
    """
    Validates a Gemini API key by sending a minimal test request.
    Returns {"valid": bool, "message": str}
    """
    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(_MODEL_NAME)
        response = model.generate_content("Reply with the single word: OK")
        if response.text and len(response.text.strip()) > 0:
            return {"valid": True, "message": "✓ Connected to Gemini 1.5 Flash successfully!"}
        return {"valid": False, "message": "API returned empty response."}
    except Exception as e:
        err = str(e)
        if "API_KEY_INVALID" in err or "API key" in err:
            return {"valid": False, "message": "✗ Invalid API key. Please check and try again."}
        return {"valid": False, "message": f"✗ Connection failed: {err[:120]}"}


def generate_cover_letter_gemini(resume_text: str, metadata: dict, job_desc: str, api_key: str) -> str:
    """
    Uses Gemini to write a tailored, professional cover letter.
    """
    import google.generativeai as genai

    genai.configure(api_key=api_key)

    skills_str = ", ".join(metadata.get("skills", [])[:6]) if metadata else ""
    prompt = f"""
You are a professional career coach and expert cover letter writer.
Write a compelling, tailored cover letter for the following candidate.

RESUME CONTENT:
{resume_text[:3000]}

TARGET JOB DESCRIPTION:
{job_desc if job_desc else "General professional role — write a strong general-purpose cover letter."}

GUIDELINES:
- Tone: Professional, confident, and engaging. Not generic.
- Length: 3 short paragraphs + opening/closing.
- Opening: Hook sentence referencing the role specifically.
- Para 1: Highlight 2-3 most relevant skills/achievements from the resume that match the JD.
- Para 2: Show enthusiasm for the company/role and cultural fit.
- Para 3: Call to action — invite interview, express readiness.
- Closing: "Sincerely, [Your Name]"
- Do NOT use hollow phrases like "I am writing to apply". Be direct and compelling.
- Output the cover letter text only. No labels, no JSON, no explanation.
"""

    try:
        model = genai.GenerativeModel(_MODEL_NAME)
        response = model.generate_content(prompt)
        return {"text": response.text.strip(), "quota_exceeded": False}
    except Exception as e:
        if is_quota_error(e):
            return {
                "text": generate_cover_letter_mock(metadata, job_desc),
                "quota_exceeded": True
            }
        return {
            "text": generate_cover_letter_mock(metadata, job_desc) + f"\n\n[Note: Gemini error — {str(e)[:100]}]",
            "quota_exceeded": False
        }


def generate_cover_letter_mock(metadata: dict, job_desc: str) -> str:
    """
    Generates a decent template cover letter using extracted resume metadata.
    """
    skills = metadata.get("skills", []) if metadata else []
    skills_str = ", ".join(skills[:4]) if skills else "a diverse professional skill set"
    role_hint = ""
    if job_desc:
        words = job_desc.split()
        caps = [w for w in words if w and w[0].isupper() and len(w) > 3]
        role_hint = caps[0] if caps else "this role"
    else:
        role_hint = "this opportunity"

    return f"""Dear Hiring Manager,

Having followed your organization's work closely, I was excited to come across {role_hint}. My background in {skills_str} positions me as a strong candidate who can contribute meaningfully from day one.

Throughout my career, I have consistently delivered results by combining technical depth with a collaborative approach. I am particularly proud of my ability to translate complex requirements into scalable, maintainable solutions — a skill I believe aligns closely with what your team is building.

I am confident that my experience and drive make me an excellent fit for your team. I would welcome the opportunity to discuss how my background can support your goals. I am available at your earliest convenience and look forward to connecting.

Sincerely,
[Your Name]"""

def generate_optimized_resume_gemini(text, metadata, job_desc="", api_key=""):
    """
    Uses Gemini to rewrite and optimize the resume content.
    Returns a structured dict representing the optimized resume.
    """
    import google.generativeai as genai
    import json

    genai.configure(api_key=api_key)

    prompt = f"""
You are a professional resume writer and ATS optimization expert. 
Your task is to rewrite the following resume to maximize ATS compatibility and impact for the target job.

ORIGINAL RESUME:
{text}

TARGET JOB DESCRIPTION:
{job_desc if job_desc else "General professional role (optimize for broad industry ATS standards)"}

Rewrite the resume and return ONLY valid JSON in this exact structure. No markdown fences, no explanation.
{{
    "name": "Candidate full name extracted from resume",
    "contact": "City, State  •  email@example.com  •  +1 (555) 000-0000  •  linkedin.com/in/profile",
    "summary": "2-3 sentence professional summary, rewritten for maximum impact and ATS alignment",
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name  •  City, State",
            "dates": "YYYY – YYYY",
            "bullets": [
                {{"text": "Original or improved bullet point", "type": "original"}},
                {{"text": "AI-rephrased bullet emphasizing impact and metrics", "type": "rephrased", "annotation": "Action-verb strength improved"}},
                {{"text": "New bullet injecting a missing keyword naturally", "type": "injected", "annotation": "Keyword added: [keyword]"}}
            ]
        }}
    ],
    "skills": {{
        "categories": [
            {{"name": "Technical", "items": ["Skill1", "Skill2", "Skill3"], "highlighted": ["Skill2"]}},
            {{"name": "Professional", "items": ["Skill4", "Skill5"], "highlighted": []}}
        ]
    }},
    "education": [
        {{
            "degree": "Degree Title",
            "institution": "University Name",
            "dates": "Month YYYY"
        }}
    ],
    "changes_summary": [
        "Brief description of key change #1 made",
        "Brief description of key change #2 made",
        "Brief description of key change #3 made"
    ]
}}

Rules:
- Keep ALL real information from the resume (name, employer names, dates, degree). Do NOT fabricate facts.
- Improve phrasing, add metrics where plausible, inject 2-3 missing JD keywords naturally.
- Mark bullet types: "original" (unchanged), "rephrased" (improved phrasing), "injected" (new keyword bullet).
- Limit to max 2 experience entries for brevity if resume is long.
- Return only raw JSON, no markdown.
"""

    try:
        model = genai.GenerativeModel(_MODEL_NAME)
        response = model.generate_content(prompt)
        response_text = response.text.strip()

        # Strip markdown fences if present
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        response_text = response_text.strip()

        result = json.loads(response_text)
        result["source"] = "gemini"
        return result

    except Exception as e:
        fallback = generate_optimized_resume_mock(text, metadata, job_desc)
        if is_quota_error(e):
            fallback["quota_exceeded"] = True
            fallback["error"] = "⚠ Built-in API quota reached. Please add your own Gemini API key in Settings."
        else:
            fallback["error"] = f"Gemini error: {str(e)[:100]}"
        return fallback


def generate_optimized_resume_mock(text, metadata, job_desc=""):
    """
    Returns a structured optimized resume using extracted metadata as fallback.
    """
    skills = metadata.get("skills", [])
    sections = metadata.get("sections", [])
    filename = metadata.get("filename", "resume.pdf")

    # Try to extract candidate name from filename or first line of text
    name_guess = "Your Name"
    first_line = text.strip().split("\n")[0].strip() if text.strip() else ""
    if first_line and len(first_line.split()) <= 4 and first_line[0].isupper():
        name_guess = first_line

    # Build skill categories
    tech_skills = [s for s in skills if s in [
        "Python", "JavaScript", "Java", "React", "Node.js", "SQL", "Docker",
        "AWS", "TypeScript", "Git", "HTML", "CSS", "MongoDB", "PostgreSQL",
        "Figma", "Next.js", "TensorFlow", "PyTorch", "Kubernetes"
    ]]
    soft_skills = [s for s in skills if s not in tech_skills]

    if not tech_skills:
        tech_skills = skills[:4]
        soft_skills = skills[4:8]

    highlighted_tech = tech_skills[:2]
    highlighted_soft = soft_skills[:1] if soft_skills else []

    missing_kws = ["CI/CD", "System Design", "Unit Testing", "Agile", "REST API"]
    injected_kw = missing_kws[0] if missing_kws else "API Integration"

    return {
        "source": "mock",
        "name": name_guess,
        "contact": "Your City, State  •  your.email@example.com  •  +1 (555) 000-0000",
        "summary": (
            f"Results-driven professional with demonstrated expertise in {', '.join(skills[:3]) if skills else 'technology and business'}. "
            f"Proven track record of delivering high-impact solutions aligned with organizational goals. "
            f"Skilled in cross-functional collaboration and building scalable systems."
        ),
        "experience": [
            {
                "title": "Senior Professional",
                "company": "Your Company  •  City, State",
                "dates": "2021 – Present",
                "bullets": [
                    {"text": "Led development and deployment of key product features, improving user satisfaction metrics by 35%.", "type": "rephrased", "annotation": "Action-verb + quantified impact added"},
                    {"text": f"Implemented {injected_kw} workflows, reducing delivery cycle time by 20% across engineering teams.", "type": "injected", "annotation": f"Keyword injected: {injected_kw}"},
                    {"text": "Collaborated with cross-functional teams to define product roadmaps and deliver milestones on schedule.", "type": "original"},
                ]
            },
            {
                "title": "Junior Professional",
                "company": "Previous Company  •  City, State",
                "dates": "2018 – 2021",
                "bullets": [
                    {"text": "Contributed to architecture decisions that scaled the platform to handle 10x traffic growth.", "type": "rephrased", "annotation": "Metric and scope added"},
                    {"text": "Built and maintained automated test suites, achieving 90%+ code coverage across modules.", "type": "original"},
                ]
            }
        ],
        "skills": {
            "categories": [
                {"name": "Technical", "items": tech_skills[:6] or ["Python", "SQL", "Git"], "highlighted": highlighted_tech},
                {"name": "Professional", "items": soft_skills[:5] or ["Project Management", "Agile"], "highlighted": highlighted_soft}
            ]
        },
        "education": [
            {
                "degree": "Bachelor of Science in Computer Science",
                "institution": "University Name",
                "dates": "May 2018"
            }
        ],
        "changes_summary": [
            "Rewrote 3 bullet points with stronger action verbs and quantified impact metrics.",
            f"Injected missing keyword '{injected_kw}' naturally into experience descriptions.",
            "Condensed and tightened professional summary for ATS clarity and recruiter readability."
        ]
    }


def clean_for_pdf(text: str) -> str:
    """Helper to convert common unsupported unicode characters in standard Helvetica."""
    if not text:
        return ""
    replacements = {
        "•": "|",
        "–": "-",
        "—": "-",
        "’": "'",
        "“": '"',
        "”": '"',
    }
    for orig, rep in replacements.items():
        text = text.replace(orig, rep)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def generate_optimized_resume_pdf(resume_data: dict) -> bytes:
    """
    Generates a beautifully styled, print-ready PDF document from the optimized resume.
    """
    from fpdf import FPDF
    
    pdf = FPDF(format='letter')
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Header Section
    pdf.set_font("Helvetica", "B", 20)
    name = resume_data.get("name", "Your Name")
    pdf.cell(0, 10, clean_for_pdf(name), align="C")
    pdf.ln(10)
    
    pdf.set_font("Helvetica", "", 9.5)
    contact = resume_data.get("contact", "")
    pdf.cell(0, 6, clean_for_pdf(contact), align="C")
    pdf.ln(8)
    
    # Section Header Helper
    def add_section(title):
        pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(37, 99, 235)  # #2563EB Primary Blue
        pdf.cell(0, 8, clean_for_pdf(title.upper()), align="L")
        pdf.ln(8)
        pdf.line(20, pdf.get_y(), 196, pdf.get_y())  # Underline across printable width
        pdf.set_text_color(0, 0, 0)
        pdf.ln(4)

    # 1. Professional Summary
    summary = resume_data.get("summary", "")
    if summary:
        add_section("Professional Summary")
        pdf.set_font("Helvetica", "", 9.5)
        pdf.multi_cell(0, 5.2, clean_for_pdf(summary))
        pdf.ln(4)

    # 2. Experience Section
    experience = resume_data.get("experience", [])
    if experience:
        add_section("Work Experience")
        for job in experience:
            # Job Title and Dates
            pdf.set_font("Helvetica", "B", 10)
            title = job.get("title", "")
            dates = job.get("dates", "")
            pdf.cell(130, 6, clean_for_pdf(title), align="L")
            pdf.set_font("Helvetica", "I", 9.5)
            pdf.cell(0, 6, clean_for_pdf(dates), align="R")
            pdf.ln(6)
            
            # Company Details
            pdf.set_font("Helvetica", "BI", 9.5)
            pdf.set_text_color(75, 85, 99)  # Gray
            pdf.cell(0, 5, clean_for_pdf(job.get("company", "")))
            pdf.ln(6)
            pdf.set_text_color(0, 0, 0)
            
            # Bullet Points
            pdf.set_font("Helvetica", "", 9)
            for bullet in job.get("bullets", []):
                bullet_text = bullet.get("text", "")
                pdf.set_x(25)  # Indent bullet point
                pdf.cell(5, 5.2, "-", border=0)
                # Pass explicit width (pdf.epw - 10) to prevent fpdf2 width calculation errors
                pdf.multi_cell(pdf.epw - 10, 5.2, clean_for_pdf(bullet_text))
            pdf.ln(3)

    # 3. Skills Section
    skills = resume_data.get("skills", {})
    categories = skills.get("categories", [])
    if categories:
        add_section("Skills & Expertise")
        for cat in categories:
            pdf.set_font("Helvetica", "B", 9.5)
            cat_name = cat.get("name", "") + ": "
            pdf.write(5.2, clean_for_pdf(cat_name))
            
            pdf.set_font("Helvetica", "", 9.5)
            items_str = ", ".join(cat.get("items", []))
            pdf.write(5.2, clean_for_pdf(items_str))
            pdf.ln(6)
        pdf.ln(4)

    # 4. Education Section
    education = resume_data.get("education", [])
    if education:
        add_section("Education")
        for edu in education:
            pdf.set_font("Helvetica", "B", 10)
            degree = edu.get("degree", "")
            dates = edu.get("dates", "")
            pdf.cell(130, 6, clean_for_pdf(degree), align="L")
            pdf.set_font("Helvetica", "I", 9.5)
            pdf.cell(0, 6, clean_for_pdf(dates), align="R")
            pdf.ln(6)
            
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(75, 85, 99)
            pdf.cell(0, 5, clean_for_pdf(edu.get("institution", "")))
            pdf.ln(6)
            pdf.set_text_color(0, 0, 0)

    return bytes(pdf.output())


